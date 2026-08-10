from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── Helpers ──────────────────────────────────────────────────────────────────

def add_run(para, text, bold=False, italic=False, size=10, color=None, name='Calibri'):
    run = para.add_run(text)
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run

def add_heading_styled(doc, text, level=2, color=None, size=13):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return h

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('─' * 110)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(200, 200, 200)

def add_experience_block(doc, exp, BLUE, DGRAY, MGRAY, is_first=False):
    """Render a single experience entry."""
    if not is_first:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Title | Company
    job_para = doc.add_paragraph()
    job_para.paragraph_format.space_after = Pt(0)
    add_run(job_para, exp['title'] + '  ', bold=True, size=11, color=DGRAY)
    add_run(job_para, exp['company'], bold=True, size=11, color=BLUE)

    # Period
    per_para = doc.add_paragraph(exp['period'])
    per_para.paragraph_format.space_before = Pt(1)
    per_para.paragraph_format.space_after = Pt(2)
    per_para.runs[0].font.italic = True
    per_para.runs[0].font.size = Pt(9)
    per_para.runs[0].font.color.rgb = MGRAY

    # Key achievement (highlight)
    if exp.get('achievement'):
        ach_para = doc.add_paragraph()
        ach_para.paragraph_format.space_before = Pt(0)
        ach_para.paragraph_format.space_after = Pt(3)
        ach_para.paragraph_format.left_indent = Inches(0.1)
        add_run(ach_para, '► ', bold=True, size=9.5, color=BLUE)
        add_run(ach_para, exp['achievement'], bold=False, size=9.5, color=DGRAY)

    # Bullet duties
    for duty in exp['duties']:
        b = doc.add_paragraph(style='List Bullet')
        b.paragraph_format.space_before = Pt(0)
        b.paragraph_format.space_after = Pt(1)
        add_run(b, duty, size=9.5, color=MGRAY)

    # Skills line
    sk = doc.add_paragraph()
    sk.paragraph_format.space_before = Pt(2)
    sk.paragraph_format.space_after = Pt(0)
    add_run(sk, 'Skills: ', bold=True, size=8.5, color=DGRAY)
    add_run(sk, exp['skills'], size=8.5, color=BLUE)


# ─── Main generator ───────────────────────────────────────────────────────────

def create_cv(language='greek'):
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Colours
    BLUE  = RGBColor(0x00, 0x7A, 0xCC)   # primary accent
    DGRAY = RGBColor(0x33, 0x33, 0x33)   # dark text
    MGRAY = RGBColor(0x55, 0x55, 0x55)   # body text
    GREEN = RGBColor(0x16, 0xA3, 0x4A)   # available badge

    # ── HEADER ────────────────────────────────────────────────────────────────
    if language == 'greek':
        name_text = 'ΑΘΑΝΑΣΙΟΣ ΖΗΣΟΓΛΟΥ'
        title_text = 'Embedded Software Engineer'
        contact_line = 'zisoglou@hotmail.gr  |  linkedin.com/in/athanasios-zisoglou-a3841561  |  github.com/ZiSo89'
        available_text = '● Ανοιχτός σε ευκαιρίες'
    else:
        name_text = 'ATHANASIOS ZISOGLOU'
        title_text = 'Embedded Software Engineer'
        contact_line = 'zisoglou@hotmail.gr  |  linkedin.com/in/athanasios-zisoglou-a3841561  |  github.com/ZiSo89'
        available_text = '● Open to opportunities'

    name_para = doc.add_heading(name_text, level=1)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in name_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = DGRAY

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(2)
    add_run(title_para, title_text, bold=False, size=12, color=BLUE)

    avail_para = doc.add_paragraph()
    avail_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    avail_para.paragraph_format.space_before = Pt(0)
    avail_para.paragraph_format.space_after = Pt(3)
    av_run = add_run(avail_para, available_text, bold=True, size=9, color=GREEN)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(4)
    add_run(contact_para, contact_line, size=9, color=MGRAY)

    add_separator(doc)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    if language == 'greek':
        add_heading_styled(doc, 'ΠΕΡΙΛΗΨΗ', level=2, color=BLUE, size=12)
        summary_text = (
            'Embedded Software Engineer με πάνω από 11 χρόνια εμπειρίας σε BMS firmware, '
            'C/C++ συστήματα, βιομηχανικό αυτοματισμό και διαδραστικές εγκαταστάσεις. '
            'Τελευταία ανέλαβε commissioning μεγάλης κλίμακας συστήματος αυτοματισμού '
            'αποθήκης στην Knapp AG (Ιούν. 2025 – Φεβ. 2026). Εστιάζει στην ποιότητα, '
            'την αποδοτικότητα διαδικασιών και τη συνεχή βελτίωση στην παράδοση λογισμικού.'
        )
    else:
        add_heading_styled(doc, 'PROFESSIONAL SUMMARY', level=2, color=BLUE, size=12)
        summary_text = (
            'Embedded Software Engineer with 11+ years of experience across BMS firmware, '
            'C/C++ systems, industrial automation, and interactive installations. Most recently '
            'commissioned a large-scale warehouse automation system at Knapp AG '
            '(Jun 2025 – Feb 2026). Focused on quality, process efficiency, and continuous '
            'improvement in software delivery and performance.'
        )

    s = doc.add_paragraph(summary_text)
    s.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.space_after = Pt(4)
    for run in s.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = MGRAY

    add_separator(doc)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    if language == 'greek':
        add_heading_styled(doc, 'ΕΜΠΕΙΡΙΑ', level=2, color=BLUE, size=12)
        experiences = [
            {
                'title': 'Software Commissioning Engineer',
                'company': 'Knapp AG',
                'period': 'Ιούνιος 2025 – Φεβρουάριος 2026',
                'achievement': 'Παρέδωσε software commissioning για μεγάλης κλίμακας σύστημα αυτοματισμού αποθήκης· ηγήθηκε λειτουργικών δοκιμών και εκπαίδευσης πελατών πριν από το go-live.',
                'duties': [
                    'Παραμετροποίηση και βελτιστοποίηση λογισμικού βάσει απαιτήσεων πελατών',
                    'Εκτέλεση δοκιμών με προσομοιωτές και επιτόπιες δοκιμές ολοκλήρωσης',
                    'Υποστήριξη πελατών κατά τη φάση έναρξης λειτουργίας και εκπαίδευσης',
                ],
                'skills': 'C++, Automation Systems, Integration Testing, Customer Training',
            },
            {
                'title': 'Embedded Software Engineer',
                'company': 'Sunlight Group Storage Systems',
                'period': 'Ιούνιος 2022 – Σεπτέμβριος 2024',
                'achievement': 'Παρέδωσε BMS firmware για 3 νέες σειρές προϊόντων εντός χρονοδιαγράμματος.',
                'duties': [
                    'Ανάπτυξη και δοκιμή BMS λογισμικού για νέα προϊόντα ενεργειακής αποθήκευσης',
                    'Ανάλυση και βελτίωση υφιστάμενων modules σύμφωνα με απαιτήσεις πελατών',
                    'Μείωση ελαττωμάτων post-release με εισαγωγή συστηματικής διαδικασίας peer code review',
                ],
                'skills': 'C, STM32, CANopen, IAR Embedded Workbench, RTOS',
            },
            {
                'title': 'Software Developer',
                'company': 'Interactive Displays GmbH',
                'period': 'Μάρτιος 2020 – Φεβρουάριος 2022',
                'achievement': 'Το προϊόν Mirotouch παραδόθηκε σε 10+ πελάτες σε όλη την Ευρώπη.',
                'duties': [
                    'Ανάπτυξη προϊόντων αφής με χρήση OpenCV και Qt framework',
                    'Προγραμματισμός μικροελεγκτών (Arduino, Raspberry Pi)',
                    'Ενσωμάτωση κάμερας βάθους Intel RealSense για ανίχνευση χειρονομιών',
                ],
                'skills': 'C++, Qt, OpenCV, Linux, Raspberry Pi, Intel RealSense',
            },
            {
                'title': 'Software Engineer',
                'company': 'Intralot',
                'period': 'Οκτώβριος 2019 – Φεβρουάριος 2020',
                'achievement': None,
                'duties': [
                    'Παροχή υπηρεσιών υποστήριξης 3ου επιπέδου για εφαρμογές λογισμικού',
                    'Αντιμετώπιση προβλημάτων και ανάπτυξη μόνιμων διορθώσεων λογισμικού',
                ],
                'skills': 'C, UNIX/Linux, Jira',
            },
            {
                'title': 'Software Developer',
                'company': 'Kurt Hüttinger GmbH',
                'period': 'Ιανουάριος 2017 – Απρίλιος 2019',
                'achievement': 'Συνεισέφερε σε 50+ εγκαταστάσεις μουσείων σε 8 χώρες της ΕΕ.',
                'duties': [
                    'Ανάπτυξη ενσωματωμένων και multimedia εφαρμογών με Arduino και Raspberry Pi',
                    'Υλοποίηση διαδραστικών εκθεμάτων για Μουσεία Επιστήμης & Τεχνολογίας',
                ],
                'skills': 'Embedded Systems, Sensors, LED Control, RS485, UART, SPI, I2C',
            },
            {
                'title': 'Ηλεκτρονικός & Ηλεκτρολόγος – Τμήμα Service & Εγκαταστάσεων',
                'company': 'Kurt Hüttinger GmbH',
                'period': 'Αύγουστος 2015 – Ιανουάριος 2017',
                'achievement': None,
                'duties': [
                    'Εγκατάσταση και συντήρηση συστημάτων σε Μουσεία Επιστήμης σε όλη την ΕΕ (75% ταξίδια)',
                    'Εκτέλεση ηλεκτρονικών δοκιμών, troubleshooting και αναβαθμίσεις',
                ],
                'skills': 'Siemens Logo! PLC, C, JavaScript, Hardware Integration, On-site Commissioning',
            },
            {
                'title': 'Πρακτική – Junior Software Developer',
                'company': 'Kurt Hüttinger GmbH',
                'period': 'Οκτώβριος 2014 – Μάιος 2015',
                'achievement': None,
                'duties': [
                    'Έρευνα σε διεπαφές PC–PLC και πρωτοτυποποίηση νέων λύσεων για εκθέματα μουσείων',
                    'Πρώτη διεθνής εμπειρία σε διεπιστημονική γερμανική ομάδα μηχανικών',
                ],
                'skills': 'PLC Programming, Prototyping, Technical Documentation',
            },
            {
                'title': 'Πρακτική – Junior Computer & Network Technician',
                'company': 'Prisma Electronics',
                'period': 'Απρίλιος 2014 – Αύγουστος 2014',
                'achievement': None,
                'duties': [
                    'Υποστήριξη σε ρύθμιση δικτύων, συντήρηση PC και troubleshooting συστημάτων',
                ],
                'skills': 'Networking, Hardware Setup, Troubleshooting',
            },
        ]
    else:
        add_heading_styled(doc, 'PROFESSIONAL EXPERIENCE', level=2, color=BLUE, size=12)
        experiences = [
            {
                'title': 'Software Commissioning Engineer',
                'company': 'Knapp AG',
                'period': 'June 2025 – February 2026',
                'achievement': 'Delivered software commissioning for a large-scale warehouse automation system; led functional tests and customer training before go-live.',
                'duties': [
                    'Configured and optimized software based on specific customer requirements',
                    'Executed simulation and on-site integration testing ensuring high software reliability',
                    'Supported clients during commissioning and training phases',
                ],
                'skills': 'C++, Automation Systems, Integration Testing, Customer Training',
            },
            {
                'title': 'Embedded Software Engineer',
                'company': 'Sunlight Group Storage Systems',
                'period': 'June 2022 – September 2024',
                'achievement': 'Delivered BMS firmware for 3 new product lines on schedule.',
                'duties': [
                    'Developed and tested BMS software for new energy storage products',
                    'Analyzed and extended existing modules based on client requirements',
                    'Reduced post-release defects by introducing a systematic peer code review process',
                ],
                'skills': 'C, STM32, CANopen, IAR Embedded Workbench, RTOS',
            },
            {
                'title': 'Software Developer',
                'company': 'Interactive Displays GmbH',
                'period': 'March 2020 – February 2022',
                'achievement': 'Mirotouch product shipped to 10+ clients across Europe.',
                'duties': [
                    'Designed and implemented touch-enabled systems using OpenCV and Qt',
                    'Developed microcontroller applications (Arduino, Raspberry Pi)',
                    'Integrated Intel RealSense depth cameras for gesture and touch detection',
                ],
                'skills': 'C++, Qt, OpenCV, Linux, Raspberry Pi, Intel RealSense',
            },
            {
                'title': 'Software Engineer',
                'company': 'Intralot',
                'period': 'October 2019 – February 2020',
                'achievement': None,
                'duties': [
                    'Delivered 3rd-level software support and implemented automation controls',
                    'Diagnosed and resolved software issues by deploying permanent fixes',
                ],
                'skills': 'C, UNIX/Linux, Jira',
            },
            {
                'title': 'Software Developer',
                'company': 'Kurt Hüttinger GmbH',
                'period': 'January 2017 – April 2019',
                'achievement': 'Contributed to 50+ museum installations across 8 EU countries.',
                'duties': [
                    'Developed embedded and multimedia applications using Arduino and Raspberry Pi',
                    'Built interactive exhibit software for Science & Technology Museums across Europe',
                ],
                'skills': 'Embedded Systems, Sensors, LED Control, RS485, UART, SPI, I2C',
            },
            {
                'title': 'Electronic & Electrician – Service & Installation Dept.',
                'company': 'Kurt Hüttinger GmbH',
                'period': 'August 2015 – January 2017',
                'achievement': None,
                'duties': [
                    'Installed and serviced systems for Science & Technology Museums across the EU (75% travel)',
                    'Executed electronic testing, troubleshooting, and upgrades based on client specifications',
                ],
                'skills': 'Siemens Logo! PLC, C, JavaScript, Hardware Integration, On-site Commissioning',
            },
            {
                'title': 'Internship – Junior Software Developer',
                'company': 'Kurt Hüttinger GmbH',
                'period': 'October 2014 – May 2015',
                'achievement': None,
                'duties': [
                    'Researched PC–PLC communication interfaces; prototyped new solutions for interactive museum exhibits',
                    'Gained first international experience in a cross-functional German engineering team',
                ],
                'skills': 'PLC Programming, Prototyping, Technical Documentation',
            },
            {
                'title': 'Internship – Junior Computer & Network Technician',
                'company': 'Prisma Electronics',
                'period': 'April 2014 – August 2014',
                'achievement': None,
                'duties': [
                    'Assisted in network configuration, PC maintenance, and system troubleshooting',
                ],
                'skills': 'Networking, Hardware Setup, Troubleshooting',
            },
        ]

    for i, exp in enumerate(experiences):
        add_experience_block(doc, exp, BLUE, DGRAY, MGRAY, is_first=(i == 0))

    add_separator(doc)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    if language == 'greek':
        add_heading_styled(doc, 'ΕΚΠΑΙΔΕΥΣΗ', level=2, color=BLUE, size=12)
        education = [
            {
                'degree': 'B.Sc. Μηχανικός Υπολογιστών',
                'institution': 'ΤΕΙ Ανατολικής Μακεδονίας & Θράκης, Ελλάδα',
                'period': '2007 – 2015',
                'notes': ['Τεχνητή Νοημοσύνη & Ρομποτική, Ενσωματωμένα Συστήματα, C++, PLC'],
            },
            {
                'degree': 'Erasmus Exchange Program  [Erasmus+]',
                'institution': 'Akademia Techniczno-Humanistyczna, Πολωνία',
                'period': '2011 – 2012',
                'notes': ['Εστίαση στα ενσωματωμένα συστήματα – διεθνής ακαδημαϊκή εμπειρία'],
            },
            {
                'degree': 'Συνεχής Εκπαίδευση',
                'institution': 'Coursera – Online Learning Platform',
                'period': '2023 – σήμερα',
                'notes': [
                    'Machine Learning Specialization – Andrew Ng, Stanford University',
                    'Deep Learning Specialization – DeepLearning.AI',
                    'AI For Everyone – Andrew Ng, DeepLearning.AI',
                ],
            },
        ]
    else:
        add_heading_styled(doc, 'EDUCATION', level=2, color=BLUE, size=12)
        education = [
            {
                'degree': 'B.Sc. in Computer Engineering',
                'institution': 'Eastern Macedonia & Thrace Institute of Technology, Greece',
                'period': '2007 – 2015',
                'notes': ['Artificial Intelligence & Robotics, Embedded Systems, C++, Industrial PLC Automation'],
            },
            {
                'degree': 'Erasmus Exchange Program  [Erasmus+]',
                'institution': 'Akademia Techniczno-Humanistyczna, Poland',
                'period': '2011 – 2012',
                'notes': ['Focused on embedded systems development – international academic experience'],
            },
            {
                'degree': 'Continuing Education',
                'institution': 'Coursera – Online Learning Platform',
                'period': '2023 – Present',
                'notes': [
                    'Machine Learning Specialization – Andrew Ng, Stanford University',
                    'Deep Learning Specialization – DeepLearning.AI',
                    'AI For Everyone – Andrew Ng, DeepLearning.AI',
                ],
            },
        ]

    for edu in education:
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(3)
        ep.paragraph_format.space_after = Pt(0)
        add_run(ep, edu['degree'] + '  ', bold=True, size=10.5, color=DGRAY)
        add_run(ep, edu['institution'], size=10.5, color=BLUE)
        add_run(ep, '  (' + edu['period'] + ')', italic=True, size=9, color=MGRAY)
        for note in edu['notes']:
            np_ = doc.add_paragraph(style='List Bullet')
            np_.paragraph_format.space_before = Pt(0)
            np_.paragraph_format.space_after = Pt(1)
            add_run(np_, note, size=9, color=MGRAY)

    add_separator(doc)

    # ── SKILLS ────────────────────────────────────────────────────────────────
    if language == 'greek':
        add_heading_styled(doc, 'ΤΕΧΝΙΚΕΣ ΔΕΞΙΟΤΗΤΕΣ', level=2, color=BLUE, size=12)
        tiers = [
            ('Expert',     'C, C++, STM32, IAR, Git, CANopen, FBD'),
            ('Proficient', 'JavaScript, Qt, OpenCV, RTOS, Linux, C#/.NET, Siemens SIMATIC Step 7, Soft Comfort'),
            ('Familiar',   'React.js, Node.js, MongoDB, Python'),
        ]
        learning_label = 'Μαθαίνω τώρα:'
        learning_text  = 'AI/ML — Machine Learning Specialization & Deep Learning Specialization (Coursera)'
    else:
        add_heading_styled(doc, 'TECHNICAL SKILLS', level=2, color=BLUE, size=12)
        tiers = [
            ('Expert',     'C, C++, STM32, IAR, Git, CANopen, FBD'),
            ('Proficient', 'JavaScript, Qt, OpenCV, RTOS, Linux, C#/.NET, Siemens SIMATIC Step 7, Soft Comfort'),
            ('Familiar',   'React.js, Node.js, MongoDB, Python'),
        ]
        learning_label = 'Currently learning:'
        learning_text  = 'AI/ML — Machine Learning Specialization & Deep Learning Specialization (Coursera)'

    tier_colors = [BLUE, RGBColor(0x63, 0x66, 0xF1), MGRAY]
    for (tier_name, skills_text), color in zip(tiers, tier_colors):
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(2)
        tp.paragraph_format.space_after = Pt(1)
        add_run(tp, tier_name + ':  ', bold=True, size=9.5, color=color)
        add_run(tp, skills_text, size=9.5, color=MGRAY)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(0)
    add_run(lp, learning_label + '  ', bold=True, size=9.5, color=RGBColor(0x16, 0xA3, 0x4A))
    add_run(lp, learning_text, size=9.5, color=MGRAY)

    add_separator(doc)

    # ── PROJECTS ──────────────────────────────────────────────────────────────
    if language == 'greek':
        add_heading_styled(doc, 'ΕΡΓΑ', level=2, color=BLUE, size=12)
        projects = [
            {
                'name': 'Mirotouch',
                'sub': 'Interactive Displays GmbH',
                'desc': 'Καινοτόμο προϊόν που μετατρέπει κάθε οθόνη σε επιφάνεια αφής με Intel RealSense και Raspberry Pi. Παραδόθηκε σε 10+ πελάτες στην Ευρώπη.',
                'tech': 'C++, Qt, OpenCV, Linux, Intel RealSense',
                'link': 'https://vimeo.com/453252636',
            },
            {
                'name': 'FastDelivery',
                'sub': 'Full-Stack Project  |  40+ REST API endpoints, 4 user roles, Real-time tracking via Socket.IO',
                'desc': 'Πλήρης πλατφόρμα διαχείρισης παραδόσεων: web admin panel, store dashboards, mobile apps για οδηγούς & πελάτες με live order tracking.',
                'tech': 'Node.js, Express, MongoDB, React.js, React Native, Socket.IO, Google Maps API, JWT',
                'link': 'https://github.com/ZiSo89/FastDelivery',
            },
            {
                'name': 'Automated Coffee Machine',
                'sub': 'Πτυχιακή Εργασία',
                'desc': 'Αυτόματη μηχανή καφέ που ελέγχεται μέσω Bluetooth από Android εφαρμογή.',
                'tech': 'Arduino (C++), Java, Atmega328, Bluetooth',
                'link': 'http://coffemake.wordpress.com/',
            },
        ]
    else:
        add_heading_styled(doc, 'SELECTED PROJECTS', level=2, color=BLUE, size=12)
        projects = [
            {
                'name': 'Mirotouch',
                'sub': 'Interactive Displays GmbH',
                'desc': 'Innovative product turning any screen into a touch surface using Intel RealSense depth camera and Raspberry Pi. Shipped to 10+ clients across Europe.',
                'tech': 'C++, Qt, OpenCV, Linux, Intel RealSense',
                'link': 'https://vimeo.com/453252636',
            },
            {
                'name': 'FastDelivery',
                'sub': 'Full-Stack Project  |  40+ REST API endpoints, 4 user roles, Real-time tracking via Socket.IO',
                'desc': 'Full delivery management platform: web admin panel, store dashboards, driver & customer mobile apps with live order tracking.',
                'tech': 'Node.js, Express, MongoDB, React.js, React Native, Socket.IO, Google Maps API, JWT',
                'link': 'https://github.com/ZiSo89/FastDelivery',
            },
            {
                'name': 'Automated Coffee Machine',
                'sub': 'Diploma Thesis',
                'desc': 'Fully automated coffee maker controlled via Bluetooth Android app.',
                'tech': 'Arduino (C++), Java, Atmega328, Bluetooth',
                'link': 'http://coffemake.wordpress.com/',
            },
        ]

    for proj in projects:
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after = Pt(0)
        add_run(pp, proj['name'] + '  ', bold=True, size=10.5, color=DGRAY)
        add_run(pp, proj['sub'], italic=True, size=9.5, color=BLUE)

        dp = doc.add_paragraph(proj['desc'])
        dp.paragraph_format.space_before = Pt(1)
        dp.paragraph_format.space_after = Pt(1)
        dp.runs[0].font.size = Pt(9.5)
        dp.runs[0].font.color.rgb = MGRAY
        dp.runs[0].font.name = 'Calibri'

        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after = Pt(1)
        add_run(tp, 'Tech: ', bold=True, size=8.5, color=DGRAY)
        add_run(tp, proj['tech'], size=8.5, color=BLUE)
        add_run(tp, '  —  ' + proj['link'], size=8.5, color=MGRAY)

    add_separator(doc)

    # ── LANGUAGES & INTERESTS ─────────────────────────────────────────────────
    if language == 'greek':
        add_heading_styled(doc, 'ΓΛΩΣΣΕΣ & ΕΝΔΙΑΦΕΡΟΝΤΑ', level=2, color=BLUE, size=12)
        lang_line = 'Ελληνικά (μητρική), Αγγλικά (επαγγελματική ευχέρεια), Γερμανικά (βασικά)'
        int_line  = 'Υπαίθριες δραστηριότητες (σκι, ποδηλασία βουνού), sci-fi σειρές, μαγειρική, τεχνολογίες web development'
        lang_label_text = 'Γλώσσες: '
        int_label_text  = 'Ενδιαφέροντα: '
    else:
        add_heading_styled(doc, 'LANGUAGES & INTERESTS', level=2, color=BLUE, size=12)
        lang_line = 'Greek (native), English (business fluent), German (basic)'
        int_line  = 'Outdoor activities (skiing, mountain biking), sci-fi series, cooking, latest web development trends'
        lang_label_text = 'Languages: '
        int_label_text  = 'Interests: '

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after = Pt(1)
    add_run(lp, lang_label_text, bold=True, size=10, color=DGRAY)
    add_run(lp, lang_line, size=10, color=MGRAY)

    ip = doc.add_paragraph()
    ip.paragraph_format.space_before = Pt(0)
    ip.paragraph_format.space_after = Pt(0)
    add_run(ip, int_label_text, bold=True, size=10, color=DGRAY)
    add_run(ip, int_line, size=10, color=MGRAY)

    return doc


# ─── Entry point ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("=" * 60)
    print("  CV Generator — Athanasios Zisoglou (2026)")
    print("=" * 60)

    print("\nDημιουργία Ελληνικού CV...")
    gr = create_cv('greek')
    gr.save('Zisoglou-Athanasios-CV-GR.docx')
    print("  ✓ Zisoglou-Athanasios-CV-GR.docx")

    print("\nCreating English CV...")
    en = create_cv('english')
    en.save('Zisoglou-Athanasios-CV-EN.docx')
    print("  ✓ Zisoglou-Athanasios-CV-EN.docx")

    print("\n" + "=" * 60)
    print("  Done! Both CVs saved in the files/ directory.")
    print("=" * 60)
